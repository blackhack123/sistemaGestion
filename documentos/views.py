# -*- coding: utf-8 -*-
#mportamos lib para folders
import os
from os import listdir
#para upload
from django.core.files.storage import FileSystemStorage
#IMPORTAMOS LA CONEXION PARA QUERY PERSONALIZADO
from django.db import connection
#importamos el dictfechtall
from principal.views import dictfetchall
from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
import json
from openpyxl import Workbook
#IMPORTAMOS BORDES
from openpyxl.styles import Border, Side, Alignment, Color, Fill, colors, Font
from openpyxl.cell import Cell
from django.utils.encoding import smart_str

from areas.models import Procesos
from documentos.forms import DocumentacionForm
from documentos.models import Documento, Revision
from personal.models import Personal
#importar notificacion admin
#decorador para eliminar el token en vista de file tree
from django.views.decorators.csrf import csrf_exempt
# Create your views here.


def directorios(request):
    
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
     
        mypath = 'media/gestionDocumental/'
        allfiles = listdir(mypath)
        
        context={
            'fileList':allfiles,
            'procesos':Procesos.objects.values('proceso').all().filter(estado=1),
            'personal':Personal.objects.values('id', 'nombre').all().filter(estado=1),
        }

        return render(request, 'gestionDocumental/directorios.html', context)


@csrf_exempt
def dirlist(request):
   import os
   import urllib
   r=['<ul class="jqueryFileTree" style="display: none;">']
   try:
       r=['<ul class="jqueryFileTree" style="display: none;">']
       #d=urllib.unquote(request.POST.get('dir','c:\\temp'))
       d = urllib.unquote(request.POST.get('dir'))
       for f in os.listdir(d):
           ff=os.path.join(d,f)
           if os.path.isdir(ff):
               r.append('<li class="directory collapsed"><a href="#" rel="%s/">%s</a></li>' % (ff,f))
           else:
               e=os.path.splitext(f)[1][1:] #obtener la extension
               r.append('<li class="file ext_%s"><a href="#" rel="%s">%s</a></li>' % (e,ff,f))
       r.append('</ul>')
   except Exception,e:
       r.append('No se puede cargar el Directorio: %s' % str(e))
   r.append('</ul>')
   return HttpResponse(''.join(r))


def documentos(request):


    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        """
        rootDir = 'media/gestionDocumental/'

        #LISTADO DE SOLO DIRECTORIOS
        alldirectorys = []
        for dirName, subdirList, fileList in os.walk(rootDir):
            #validamos que vaya la root
            if dirName != 'media/gestionDocumental/':
                
                element={}
                #ruta_absoluta
                element['ruta_absoluta']=dirName
                #se hace split para mostrar solo el nombre del folder no completo
                element['nombre']= dirName.split('media/gestionDocumental/')
                alldirectorys.append(element)
       
        #LISTAR DIRECTORIOS Y ARCHIVOS
        list_dir=[]
        for dirName, subdirList, fileList in os.walk(rootDir):
            element={}
            
            #validamos que no vaya la carpeta root
            if dirName != 'media/gestionDocumental/':
                #nombre original
                element['carpeta'] = dirName.split(rootDir)
                #url/ruta
                element['url_carpeta'] = dirName
                
                #archivos
                element['archivos'] = fileList
                #agregamos element al array
                list_dir.append(element)
        """
        context={
            #'directoriosList':alldirectorys,
            #'prueba':list_dir,
            'documentoForm':DocumentacionForm(),
        }

        return render(request, 'gestionDocumental/documentos.html', context)



def createDirectorio(request):


    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        if request.method == 'POST' :
            #nombre del directorio
            directorio=request.POST.get('directorio')
            #destino
            destino = request.POST.get('destino')
            try:
                os.makedirs(destino+"/"+directorio, mode=0777)
                return redirect('documentos')
            except OSError:
                pass
                return redirect('documentos')




def deleteDirectorio(request):
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        if request.method == 'POST' :
            
            #ruta del directorio
            url_carpeta=request.POST.get('url_carpeta')

            try:
                #eliminar dictorio y su contenido
                import shutil
                shutil.rmtree(url_carpeta)

                #return success
                data = {
                    'resultado':'ok_select',
                }
                return HttpResponse(json.dumps(data), content_type="application/json")
            except OSError:
                pass
                return redirect('documentos')

                


def editDirectorio(request):

    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        if request.method == 'POST' :
            #nombre anterio
            nombreAnterior=str(request.POST.get('nombreAnterior'))
            #nuevo nombre
            nuevoNombre=str(request.POST.get('nuevonombre'))
            try:
                os.rename("media/gestionDocumental/"+nombreAnterior, "media/gestionDocumental/"+nuevoNombre)
                
                return redirect('documentos')
            except OSError:
                pass
                return redirect('documentos')

def uploadArchivo(request):
  

    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        form = DocumentacionForm(request.FILES)
        #if request.method == 'POST' and request.FILES:
        #archivo
        myfile = request.FILES['archivo']
        #carpeta seleccionada
        destino=request.POST.get('destino')
        #destino final
        #rutaAbsoluta="media/gestionDocumental/"+destino+"/"
        rutaAbsoluta=destino+"/"
        fs = FileSystemStorage(location=rutaAbsoluta)
        filename = fs.save(myfile.name, myfile)
        uploaded_file_url = rutaAbsoluta+filename
        return redirect('documentos')


def selectArchivo(request):
    
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        if request.method == 'POST' and request.is_ajax():

            name = request.POST.get('archivo')
            path="media/gestionDocumental/"

            #recorrer el directorio
            for root, dirs, files in os.walk(path):

                #verificar si el archivo existe en el path
                if name.encode('utf-8') in files :
                    #ruta absoluta del directorio
                    ruta= os.path.join(root, name)

            data = {
                'resultado':'ok_select',
                'url': ruta,

            }
            return HttpResponse(json.dumps(data), content_type="application/json")



def deleteArchivo(request):

    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        if request.method == 'POST' and request.is_ajax():

            path_documento = request.POST.get('path_documento')
            """
            path="media/gestionDocumental/"
            #recorrer el directorio
            for root, dirs, files in os.walk(path):

                #verificar si el archivo existe en el path
                if archivo.encode('utf-8') in files :
                    #ruta absoluta del directorio
                    ruta= os.path.join(root, archivo)
                    os.remove(ruta)
            """
            import os
            os.remove(path_documento.encode('utf-8'))
            data = {
                'resultado':'ok_select',
            }
            return HttpResponse(json.dumps(data), content_type="application/json")


def designarDirectorio(request):

    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return redirect('logout')

    else:
        if request.method == 'POST' :
            id_personal=request.POST.get('id_personal')
            directorio=request.POST.get('directorio')

            Personal.objects.filter(pk=id_personal).update(
                carpeta=directorio
            )
            return redirect('directorios')


def estadoDocAdmin(request):
    
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        #return render(request, 'seguridad/login.html', {})
        return redirect('login')
    else:
        """
        cursor = connection.cursor()
        query = " SELECT DISTINCT personal_personal.nombre AS jefe_proceso, areas_procesos.proceso, areas_procesos.estado_colaborador, areas_procesos.estado_encargado, areas_procesos.estado_lider, areas_procesos.estado_admin, "
        query = query + " areas_procesos.id AS id_proceso "
        query = query + " FROM areas_procesos "
        query = query + " LEFT JOIN areas_area_proceso ON areas_procesos.id = areas_area_proceso.proceso_id "
        query = query + " LEFT JOIN areas_areas ON areas_area_proceso.area_id = areas_areas.id "
        query = query + " LEFT JOIN personal_personal ON areas_areas.id_personal_id = personal_personal.id "
 
        
        cursor.execute(query)
        
        rows = dictfetchall(cursor)
        list_estado_documentos = list(rows)
        """
        context={
            #'estado_documentos':list_estado_documentos,
            'documentos':Documento.objects.values('id', 'nombre', 'version').all().filter(estado=4)
        }
        return render(request, 'gestionDocumental/estadoDocumentacion.html', context)


def gridEstadoDocumentacion(request):
    
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return render(request, 'seguridad/login.html', {})

    else:

        cursor = connection.cursor()
        query = " SELECT DISTINCT documentos_documento.id, documentos_documento.path, documentos_documento.descripcion, "
        query = query + " documentos_documento.version, documentos_documento.estado, documentos_documento.fec_subido, "
        query = query + " documentos_documento.nombre, areas_areas.area AS proceso,  areas_procesos.proceso AS procedimiento, "
        query = query + " personal_personal.nombre AS nombre_responsable "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN areas_areas ON documentos_documento.proceso_id = areas_areas.id "
        query = query + " LEFT JOIN areas_procesos ON documentos_documento.procedimiento_id = areas_procesos.id "
        query = query + " LEFT JOIN personal_personal ON documentos_documento.subido_por_id = personal_personal.id "

        cursor.execute(query)
        rows = dictfetchall(cursor)
        
        if rows:
            #RETORNA  
            return JsonResponse(rows, safe=False)
        else:
            #RETORNA MENSAJE DE NO EXISTEN
            json = {
            'resultado':'no_ok',
            'mensaje': 'no existen Areas',
            }
            return JsonResponse(json, safe=False)
    


def gridRevisionDocumentacion(request):
    
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        return render(request, 'seguridad/login.html', {})

    else:

        id_documento = request.GET['id_documento']
        
        cursor = connection.cursor()
        query = " SELECT documentos_revision.id, documentos_revision.fec_rev_director, documentos_revision.estado_rev_director, "
        query = query + " documentos_documento.id AS id_documento, documentos_revision.observacion_rev_director, documentos_revision.fec_rev_lider, documentos_revision.estado_rev_lider, "
        query = query + " documentos_revision.observacion_rev_lider, documentos_revision.fec_rev_admin, documentos_revision.estado_rev_admin, "
        query = query + " documentos_revision.observacion_rev_admin, personal_personal.nombre AS nombre_director, lider.nombre AS nombre_lider, "
        query = query + " admin.nombre AS nombre_admin, documentos_revision.director_id, documentos_revision.lider_id, documentos_revision.admin_id "
        query = query + " FROM documentos_revision "
        query = query + " LEFT JOIN documentos_documento ON documentos_revision.documento_id = documentos_documento.id "
        query = query + " LEFT JOIN personal_personal ON documentos_revision.director_id = personal_personal.id "
        query = query + " LEFT JOIN personal_personal lider ON documentos_revision.lider_id = lider.id "
        query = query + " LEFT JOIN personal_personal admin ON documentos_revision.admin_id = admin.id "
        query = query + " WHERE documentos_revision.documento_id=%s ORDER BY id DESC "

        params = [id_documento]

        cursor.execute(query, params)
        rows = dictfetchall(cursor)
        
        if rows:
            #RETORNA  
            return JsonResponse(rows, safe=False)
        else:
            #RETORNA MENSAJE DE NO EXISTEN
            json = {
                'resultado':'no_ok',
                'mensaje': 'no existen Revisiones',
            }
            return JsonResponse(json, safe=False)



def updateDocAdmin(request):
    
    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        #return render(request, 'seguridad/login.html', {})
        return redirect('login')
    else:
        
        if request.method=="POST":

            #data desde el form
            id_documento = request.POST.get('id_documento')
            id_revision = request.POST.get('id_revision')

            estado = request.POST.get('estado_documento')


            #existe observacion
            if estado == '2':

                observacion = request.POST.get('observaciones')
            else:
                observacion = None

            Documento.objects.filter(pk=id_documento).update(
                estado = estado
            )

            import datetime
            fecha_actual = datetime.datetime.now().strftime ("%Y-%m-%d")

            if estado == '4':
                estado_admin = 3
            else:
                estado_admin = 2
            Revision.objects.filter(pk=id_revision).update(
                fec_rev_admin = fecha_actual,
                estado_rev_admin = estado_admin,
                observacion_rev_admin = observacion
            )

            #borrar notificaciones
            if 'notificacionDocumentosAdmin' in request.session:
                del request.session['notificacionDocumentosAdmin']

            #redige a la vista principal estado documentos
            return redirect('estadoDocAdmin')



def GridDocumentosAdministrador(request):


    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        #return render(request, 'seguridad/login.html', {})
        return redirect('login')
    else:
        
        cursor = connection.cursor()

        query = " SELECT documentos_documento.id, documentos_documento.nombre, documentos_documento.descripcion, "
        query = query + " documentos_documento.version, documentos_documento.fec_subido, documentos_documento.path, "
        query = query + " documentos_documento.estado, personal_personal.nombre AS responsable, "
        query = query + " areas_areas.area AS proceso, areas_procesos.proceso AS procedimiento "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN personal_personal ON documentos_documento.subido_por_id = personal_personal.id "
        query = query + " LEFT JOIN areas_areas ON documentos_documento.proceso_id = areas_areas.id "
        query = query + " LEFT JOIN areas_procesos ON documentos_documento.procedimiento_id = areas_procesos.id ORDER BY proceso "

        cursor.execute(query)

        rows = dictfetchall(cursor)

        if rows:
            #RETORNA  
            return JsonResponse(rows, safe=False)
        else:
            #RETORNA MENSAJE DE NO EXISTEN
            json = {
                'resultado':'no_ok',
                'mensaje': 'no existen Documentos',
            }
            return JsonResponse(json, safe=False)


def ReporteHistorico(request):
    documento_id = request.GET['documento_id']
    documento_singular = Revision.objects.all().filter(documento_id=documento_id).last()
    documento_list = Revision.objects.all().filter(documento_id=documento_id)

    # ESTABLECER BORDES
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    # Creamos el libro de trabajo
    wb = Workbook()

    # Definimos como nuestra hoja de trabajo, la hoja activa, por defecto la primera del libro
    ws = wb.active
    ws.title = unicode("Histórico")
    # En la celda B1 ponemos el texto 'REPORTE DE PERSONAS'
    ws.merge_cells('B2:N2')
    ws['B2'] = 'Historial de cambios'
    ws['B2'].alignment = Alignment(horizontal='center')
    ws.cell(row=2, column=2).border = thin_border
    b2 = ws['B2']
    b2.font = Font(bold=True, color=colors.DARKBLUE, size=12)
    # Creamos los encabezados desde la celda B3 hasta la E3
    ws['B3'] = 'Procedimiento: '
    ws.cell(row=3, column=2).border = thin_border
    ws.cell(row=3, column=2).alignment = Alignment(horizontal='right')
    b3 = ws['B3']
    b3.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['C3'] = smart_str(documento_singular.documento.nombre)
    ws.cell(row=3, column=3).border = thin_border
    ws.cell(row=3, column=3).alignment = Alignment(horizontal='left')
    c3 = ws['C3']
    c3.font = Font(bold=True, color=colors.BLACK, size=12)

    ws['D3'] = 'Versión: '
    ws.cell(row=3, column=4).border = thin_border
    ws.cell(row=3, column=4).alignment = Alignment(horizontal='right')
    d3 = ws['D3']
    d3.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['E3'] = documento_singular.documento.version
    ws.cell(row=3, column=5).border = thin_border
    ws.cell(row=3, column=5).alignment = Alignment(horizontal='left')
    e3 = ws['E3']
    e3.font = Font(bold=True, color=colors.BLACK, size=12)

    ws['F3'] = 'Responsable: '
    ws.cell(row=3, column=6).border = thin_border
    ws.cell(row=3, column=6).alignment = Alignment(horizontal='right')
    f3 = ws['F3']
    f3.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['G3'] = smart_str(documento_singular.documento.subido_por.nombre)
    ws.cell(row=3, column=7).border = thin_border
    ws.cell(row=3, column=7).alignment = Alignment(horizontal='left')
    g3 = ws['G3']
    g3.font = Font(bold=True, color=colors.BLACK, size=12)

    ws['B4'] = 'Descripcion: '
    ws.cell(row=4, column=2).border = thin_border
    ws.cell(row=4, column=2).alignment = Alignment(horizontal='right')
    b4 = ws['B4']
    b4.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['C4'] = documento_singular.documento.descripcion
    ws.cell(row=4, column=3).border = thin_border
    ws.cell(row=4, column=3).alignment = Alignment(horizontal='left')
    c4 = ws['C4']
    c4.font = Font(bold=True, color=colors.BLACK, size=12)

    ws['D4'] = 'Fecha Subida: '
    ws.cell(row=4, column=4).border = thin_border
    ws.cell(row=4, column=4).alignment = Alignment(horizontal='right')
    d4 = ws['D4']
    d4.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['E4'] = documento_singular.documento.fec_subido
    ws.cell(row=4, column=5).border = thin_border
    ws.cell(row=4, column=5).alignment = Alignment(horizontal='left')
    e4 = ws['E4']
    e4.font = Font(bold=True, color=colors.BLACK, size=12)

    ws['B5'] = '#'
    ws.cell(row=5, column=2).border = thin_border
    ws.cell(row=5, column=2).alignment = Alignment(horizontal='center')
    b5 = ws['B5']
    b5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['C5'] = 'Director de Area'
    ws.cell(row=5, column=3).border = thin_border
    ws.cell(row=5, column=3).alignment = Alignment(horizontal='center')
    c5 = ws['C5']
    c5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['D5'] = 'Estado'
    ws.cell(row=5, column=4).border = thin_border
    ws.cell(row=5, column=4).alignment = Alignment(horizontal='center')
    d5 = ws['D5']
    d5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['E5'] = 'Observacion'
    ws.cell(row=5, column=5).border = thin_border
    ws.cell(row=5, column=5).alignment = Alignment(horizontal='center')
    e5 = ws['E5']
    e5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['F5'] = 'Fecha revision'
    ws.cell(row=5, column=6).border = thin_border
    ws.cell(row=5, column=6).alignment = Alignment(horizontal='center')
    f5 = ws['F5']
    f5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['G5'] = 'Lider de Norma'
    ws.cell(row=5, column=7).border = thin_border
    ws.cell(row=5, column=7).alignment = Alignment(horizontal='center')
    g5 = ws['G5']
    g5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['H5'] = 'Estado'
    ws.cell(row=5, column=8).border = thin_border
    ws.cell(row=5, column=8).alignment = Alignment(horizontal='center')
    h5 = ws['H5']
    h5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['I5'] = 'Observacion'
    ws.cell(row=5, column=9).border = thin_border
    ws.cell(row=5, column=9).alignment = Alignment(horizontal='center')
    i5 = ws['I5']
    i5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['J5'] = 'Fecha revision'
    ws.cell(row=5, column=10).border = thin_border
    ws.cell(row=5, column=10).alignment = Alignment(horizontal='center')
    j5 = ws['J5']
    j5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['K5'] = 'Administrador'
    ws.cell(row=5, column=11).border = thin_border
    ws.cell(row=5, column=11).alignment = Alignment(horizontal='center')
    k5 = ws['K5']
    k5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['L5'] = 'Estado'
    ws.cell(row=5, column=12).border = thin_border
    ws.cell(row=5, column=12).alignment = Alignment(horizontal='center')
    l5 = ws['L5']
    l5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['M5'] = 'Observacion'
    ws.cell(row=5, column=13).border = thin_border
    ws.cell(row=5, column=13).alignment = Alignment(horizontal='center')
    m5 = ws['M5']
    m5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    ws['N5'] = 'Fecha revision'
    ws.cell(row=5, column=14).border = thin_border
    ws.cell(row=5, column=14).alignment = Alignment(horizontal='center')
    n5 = ws['N5']
    n5.font = Font(bold=True, color=colors.DARKBLUE, size=12)

    cont = 6
    indice = 1
    for documento in documento_list:

        ws.cell(row=cont, column=2).value = indice
        ws.cell(row=cont, column=2).border = thin_border
        ws.cell(row=cont, column=2).alignment = Alignment(horizontal='center')

        ws.cell(row=cont, column=3).value = smart_str(documento.director.nombre)
        ws.cell(row=cont, column=3).border = thin_border
        ws.cell(row=cont, column=3).alignment = Alignment(horizontal='center')

        if documento.estado_rev_director == 1 or documento.estado_rev_director  == 0:
            ws.cell(row=cont, column=4).value = 'Pendiente'
        elif documento.estado_rev_director == 2:
            ws.cell(row=cont, column=4).value = 'Observacion'
        elif documento.estado_rev_director == 3:
            ws.cell(row=cont, column=4).value = 'Aprobado'
        ws.cell(row=cont, column=4).border = thin_border
        ws.cell(row=cont, column=4).alignment = Alignment(horizontal='center')

        if documento.observacion_rev_director:
            ws.cell(row=cont, column=5).value = smart_str(documento.observacion_rev_director)
        else:
            ws.cell(row=cont, column=5).value = ''
        ws.cell(row=cont, column=5).border = thin_border
        ws.cell(row=cont, column=5).alignment = Alignment(horizontal='center')

        ws.cell(row=cont, column=6).value = documento.fec_rev_director
        ws.cell(row=cont, column=6).border = thin_border
        ws.cell(row=cont, column=6).alignment = Alignment(horizontal='center')

        ws.cell(row=cont, column=7).value = smart_str(documento.lider.nombre)
        ws.cell(row=cont, column=7).border = thin_border
        ws.cell(row=cont, column=7).alignment = Alignment(horizontal='center')

        if documento.estado_rev_lider == 1 or documento.estado_rev_lider == 0:
            ws.cell(row=cont, column=8).value = 'Pendiente'
        elif documento.estado_rev_lider == 2:
            ws.cell(row=cont, column=8).value = 'Observacion'
        elif documento.estado_rev_lider == 3:
            ws.cell(row=cont, column=8).value = 'Aprobado'
        ws.cell(row=cont, column=8).border = thin_border
        ws.cell(row=cont, column=8).alignment = Alignment(horizontal='center')

        if documento.observacion_rev_lider:
            ws.cell(row=cont, column=9).value = smart_str(documento.observacion_rev_lider)
        else:
            ws.cell(row=cont, column=9).value = ''
        ws.cell(row=cont, column=9).border = thin_border
        ws.cell(row=cont, column=9).alignment = Alignment(horizontal='center')

        ws.cell(row=cont, column=10).value = documento.fec_rev_lider
        ws.cell(row=cont, column=10).border = thin_border
        ws.cell(row=cont, column=10).alignment = Alignment(horizontal='center')

        ws.cell(row=cont, column=11).value = smart_str(documento.admin.nombre)
        ws.cell(row=cont, column=11).border = thin_border
        ws.cell(row=cont, column=11).alignment = Alignment(horizontal='center')

        if documento.estado_rev_admin == 1 or documento.estado_rev_admin == 0:
            ws.cell(row=cont, column=12).value = 'Pendiente'
        elif documento.estado_rev_admin == 2:
            ws.cell(row=cont, column=12).value = 'Observacion'
        elif documento.estado_rev_admin == 3:
            ws.cell(row=cont, column=12).value = 'Aprobado'
        ws.cell(row=cont, column=12).border = thin_border
        ws.cell(row=cont, column=12).alignment = Alignment(horizontal='center')

        if documento.observacion_rev_admin:
            ws.cell(row=cont, column=13).value = smart_str(documento.observacion_rev_admin)
        else:
            ws.cell(row=cont, column=13).value = ''
        ws.cell(row=cont, column=13).border = thin_border
        ws.cell(row=cont, column=13).alignment = Alignment(horizontal='center')

        ws.cell(row=cont, column=14).value = documento.fec_rev_admin
        ws.cell(row=cont, column=14).border = thin_border
        ws.cell(row=cont, column=14).alignment = Alignment(horizontal='center')

        cont = cont + 1
        indice = indice + 1

    # Establecemos el nombre del archivo
    nombre_archivo = "historico.xlsx"

    # Definimos que el tipo de respuesta a devolver es un archivo de microsoft excel
    response = HttpResponse(content_type="application/ms-excel")
    contenido = "attachment; filename={0}".format(nombre_archivo)
    response["Content-Disposition"] = contenido

    # ESTABLECER DIMENSIONES A COLUMNAS
    ws.column_dimensions["B"].width = 20.0
    ws.column_dimensions["C"].width = 30.0
    ws.column_dimensions["D"].width = 20.0
    ws.column_dimensions["E"].width = 30.0
    ws.column_dimensions["F"].width = 30.0
    ws.column_dimensions["G"].width = 30.0
    ws.column_dimensions["H"].width = 20.0
    ws.column_dimensions["I"].width = 20.0
    ws.column_dimensions["J"].width = 20.0
    ws.column_dimensions["K"].width = 20.0
    ws.column_dimensions["L"].width = 20.0
    ws.column_dimensions["M"].width = 20.0
    ws.column_dimensions["N"].width = 20.0

    wb.save(response)
    return response


def EstablecerVigente(request):


    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        #return render(request, 'seguridad/login.html', {})
        return redirect('login')
    else:

        id_documento=request.POST.get('id_documento')

        Documento.objects.filter(pk=id_documento).update(
            estado = 4
        )

        json = {
                'resultado':'ok_update',
                'mensaje': 'Documento Vigente !!',
        }
        return JsonResponse(json, safe=False)


def EstablecerNoVigente(request):


    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        #return render(request, 'seguridad/login.html', {})
        return redirect('login')
    else:

        id_documento=request.POST.get('id_documento')

        Documento.objects.filter(pk=id_documento).update(
            estado = 5
        )

        json = {
                'resultado':'ok_update',
                'mensaje': 'Documento No Vigente !!',
        }
        return JsonResponse(json, safe=False)



def FirmarDocumento(request):


    #validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:

        #return render(request, 'seguridad/login.html', {})
        return redirect('login')
    else:

        id_documento=request.POST.get('id_documento')

        documento_data = Documento.objects.get(pk=id_documento)

        #fechas
        cursor = connection.cursor()
        query = " SELECT documentos_revision.id, documentos_revision.fec_rev_director, documentos_revision.fec_rev_lider, "
        query = query + " documentos_revision.fec_rev_admin "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN documentos_revision ON documentos_documento.id = documentos_revision.id  "
        #query = query + " WHERE documentos_documento.id =%s "
	query = query + " WHERE documentos_revision.documento_id=%s "
        query = query + " AND "
        query = query + " documentos_revision.estado_rev_director = 3 "
        query = query + " AND documentos_revision.estado_rev_lider = 3 "
        query = query + " AND documentos_revision.estado_rev_admin = 3 "
        query = query + " "
	print query
        params = [id_documento]
        cursor.execute(query, params)
        row_fechas = dictfetchall(cursor)

        #validacion colaborador
        cursor = connection.cursor()
        query = " SELECT personal_personal.nombre AS colaborador, personal_personal.firma "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN personal_personal ON documentos_documento.subido_por_id = personal_personal.id "
        query = query + " WHERE documentos_documento.id =%s AND firma IS NOT NULL "
        params =[id_documento]

        cursor.execute(query, params)
        row_colaborador= dictfetchall(cursor)
        colaborador_list = list(row_colaborador)
        total_colaborador = len(colaborador_list)

        if total_colaborador == 0:

            data ={
                'resultado':'no_ok_colaborador',
                'mensaje':'No se puede firmar, debe subir la firma del Colaborador. '
            }
            return JsonResponse(data, safe=False)

        #validacion administrador
        cursor = connection.cursor()
        query = " SELECT personal_personal.nombre AS administrador, personal_personal.firma "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN documentos_revision ON documentos_documento.id = documentos_revision.documento_id "
        query = query + " LEFT JOIN personal_personal ON documentos_revision.admin_id = personal_personal.id "
        query = query + " WHERE documentos_documento.id =%s AND firma IS NOT NULL "
        params =[id_documento]

        cursor.execute(query, params)
        row_administrador = dictfetchall(cursor)
        admin_list = list(row_administrador)
        total_admin = len(admin_list)

        if total_admin == 0:

            data ={
                'resultado':'no_ok_admin',
                'mensaje':'No se puede firmar, debe subir la firma del Administrador. '
            }
            return JsonResponse(data, safe=False)


        # validacion director de area
        cursor = connection.cursor()
        query = " SELECT personal_personal.nombre AS director, personal_personal.firma "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN documentos_revision ON documentos_documento.id = documentos_revision.documento_id "
        query = query + " LEFT JOIN personal_personal ON documentos_revision.director_id = personal_personal.id "
        query = query + " WHERE documentos_documento.id =%s AND firma IS NOT NULL "
        params =[id_documento]

        cursor.execute(query, params)
        row_director = dictfetchall(cursor)
        director_list = list(row_director)
        total_director = len(director_list)

        if total_director == 0:

            data ={
                'resultado':'no_ok_director',
                'mensaje':'No se puede firmar, debe subir la firma del Director de Area.'
            }
            return JsonResponse(data, safe=False)

        # validacion lider de norma
        cursor = connection.cursor()
        query = " SELECT personal_personal.nombre AS lider, personal_personal.firma "
        query = query + " FROM documentos_documento "
        query = query + " LEFT JOIN documentos_revision ON documentos_documento.id = documentos_revision.documento_id "
        query = query + " LEFT JOIN personal_personal ON documentos_revision.lider_id = personal_personal.id "
        query = query + " WHERE documentos_documento.id =%s AND firma IS NOT NULL "
        params =[id_documento]

        cursor.execute(query, params)
        row_lider = dictfetchall(cursor)
        lider_list = list(row_lider)
        total_lider = len(lider_list)

        if total_lider == 0:

            data ={
                'resultado':'no_ok_lider',
                'mensaje':'No se puede firmar, debe subir la firma del Lider de Norma.'
            }
            return JsonResponse(data, safe=False)

        #armando la data para la firma
        #documento a firmar
        path_documento      = str(documento_data.path)

        #colaborador

        colaborador         = row_colaborador[0]['colaborador']
        firma_colaborador   = row_colaborador[0]['firma']

        #firma administrador
        firma_administrador = row_administrador[0]['firma']

        #firma director de area
        firma_director      = row_director[0]['firma']

        #firma lider de norma
        firma_lider      = row_lider[0]['firma']

        import docx
        from docx import *
        from docx import Document
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Mm, Cm, Pt

        """
        tpl = DocxTemplate(path_documento)

        subdoc = tpl.new_subdoc()
        subdoc.add_picture(firma_colaborador, width=Cm(1))

        context = {
            'mysubdoc': subdoc,
        }
        tpl.render(context)
        tpl.save('footer.docx')
        """

        #abrir doc
        #document = Document(path_documento)
        #p = document.add_paragraph()
        #r = p.add_run()
        #r.add_picture(firma_administrador)
        #guardar archivo
        #document.save()

        document = Document(path_documento)

        table = document.add_table(rows=3, cols=4)
        # fila 1
        row_1 = table.rows[0]
        row_1.cells[0].text = '{{subido_por}}'
        row_1.cells[1].text = '{{director}}'
        row_1.cells[2].text = '{{lider}}'
        row_1.cells[3].text = '{{admin}}'
        #fila 2
        row_2 = table.rows[1]
        row_2.cells[0].text = '{{fec_subido}}'
        row_2.cells[1].text = '{{fec_director}}'
        row_2.cells[2].text = '{{fec_lider}}'
        row_2.cells[3].text = '{{fec_admin}}'
        #fila 3
        row_3 = table.rows[2]
        row_3.cells[0].text = '{{version}}'
        #firma Director
        paragraph = row_3.cells[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(firma_director, width=200000, height=200000)
        #firma Lider
        paragraph = row_3.cells[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(firma_lider, width=200000, height=200000)
        #firma Administrador
        paragraph = row_3.cells[3].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(firma_administrador, width=200000, height=200000)

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(7)

        #guardando el document
        document.save(path_documento)


        #template doc
        #tpl = DocxTemplate('media/plantilla.docx')
        tpl = DocxTemplate(path_documento)
	print row_fechas
        fec_subido      = str(documento_data.fec_subido)
        fec_director    = str(row_fechas[0]['fec_rev_director'])
        fec_lider       = str(row_fechas[0]['fec_rev_lider'])
        fec_admin       = str(row_fechas[0]['fec_rev_admin'])

        context = {
            'subido_por': 'Redactado por: '+row_colaborador[0]['colaborador'],
            'director': 'Aprobado por: ' + row_director[0]['director'],
            'lider': 'Aprobado por: ' + row_lider[0]['lider'],
            'admin': 'Aprobado por: ' + row_administrador[0]['administrador'],
            'fec_subido': 'Fecha subido: '+fec_subido,
            'version': 'Version: '+documento_data.version,
            'fec_director': 'Fecha aprobado: ' + fec_director,
            'fec_lider': 'Fecha aprobado: ' + fec_lider,
            'fec_admin': 'Fecha aprobado: ' + fec_admin
        }

        #reemplazo imagen director
        #tpl.replace_media('media/director.jpg', firma_director)
        #tpl.replace_media('media/lider.jpg', firma_lider)
        #tpl.replace_media('media/admin.jpg', firma_administrador)
        tpl.render(context)
        #tpl.save('media/plantilla.docx')
        tpl.save(path_documento)



        data = {
            'resultado': 'ok_firma',
            'mensaje': ' Documento firmado con exito !!'
        }
        return JsonResponse(data, safe=False)



def SeleccionarDocProcedimiento(request):

    # validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:
        return redirect('logout')
    else:
        if request.method == 'POST' and request.is_ajax():

            #id documento
            id_documento = request.POST.get('id_documento')

            #documento a modificar
            documento = Documento.objects.get(pk=id_documento)

            #data
            data = {
                'resultado': 'ok_select',
                'nombre':documento.nombre,
                'descripcion': documento.descripcion,
                'version': documento.version
            }

            #retornar data
            return HttpResponse(json.dumps(data), content_type="application/json")


def ActualizarDocProcedimiento(request):

    # validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:
        return redirect('logout')
    else:
        if request.method == 'POST' and request.is_ajax():

            #data
            id_documento = request.POST.get('id_documento')
            nombre = request.POST.get('nombre')
            descripcion = request.POST.get('descripcion')
            version = request.POST.get('version')

            Documento.objects.filter(pk=id_documento).update(
                nombre = nombre,
                descripcion = descripcion,
                version = version
            )

            #data
            data = {
                'resultado': 'ok_update',
            }
            #retornar data
            return HttpResponse(json.dumps(data), content_type="application/json")



def EliminarDocProcedimiento(request):

    # validar si existe usuario logeado
    if 'nombreUsuario' not in request.session:
        return redirect('logout')
    else:
        if request.method == 'POST' and request.is_ajax():

            #id documento
            id_documento = request.POST.get('id_documento')

            documento = Documento.objects.get(pk=id_documento)
            path_documento = documento.path

            #eliminar revision
            Revision.objects.filter(documento_id=id_documento).delete()

            #eliminar documento
            Documento.objects.filter(pk=id_documento).delete()

            #eliminar documento
            if path_documento:
                try:
                    import os
                    os.remove(str(path_documento))
                except OSError:
                    pass

            # data
            data = {
                'resultado': 'ok_delete',
            }
            # retornar data
            return HttpResponse(json.dumps(data), content_type="application/json")
